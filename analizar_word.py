import docx
import openai
import pandas as pd

# Configurar la clave de API de OpenAI
openai.api_key = "apigpt"

# Cargar el archivo Word
doc = docx.Document("Actualización Lineamientos Proyecto Arbusta.docx")

# Obtener un resumen del contenido
resumen = f"El documento Word 'Actualización Lineamientos Proyecto Arbusta.docx' contiene {len(doc.paragraphs)} párrafos."

# Obtener los primeros párrafos como muestra
muestra = [p.text for p in doc.paragraphs[:5] if p.text.strip()]

# Convertir la muestra a un DataFrame de pandas para una mejor visualización
df_muestra = pd.DataFrame(muestra, columns=["Contenido"])

# Crear un prompt para GPT
prompt = f"""
Analiza la siguiente información de un documento Word:

{resumen}

Muestra de los primeros párrafos:
{df_muestra.to_string()}

Basándote en esta información, describe de qué trata este documento Word y qué tipo de contenido contiene.
"""

# Llamar a la API de OpenAI
response = openai.ChatCompletion.create(
    model="gpt-3.5-turbo",
    messages=[
        {"role": "system", "content": "Eres un asistente experto en análisis de documentos."},
        {"role": "user", "content": prompt}
    ]
)

# Imprimir la respuesta
print(response.choices[0].message['content'])
